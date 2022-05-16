from __future__ import print_function
from docx import Document

from mailmerge import MailMerge
from datetime import date

def mailmergeDoc(years,studentObject):
    template_1 = "dwight-transcript.docx"
    

    # Show a simple example
    document_1 = MailMerge(template_1)
    print("Fields included in {}: {}".format(template_1,
                                         document_1.get_merge_fields()))
    # Merge in the values
    student= {
    'firstName' :studentObject['first_name'],
    'secondName':studentObject['last_name'],
    }


    transcript = [{
    'prod_desc': 'Red Shoes',
    'price': '$10.00',
    'quantity': '2500',
    'total_purchases': '$25,000.00'
    }, {
    'prod_desc': 'Green Shirt',
    'price': '$20.00',
    'quantity': '10000',
    'total_purchases': '$200,000.00'
    }, {
    'prod_desc': 'Purple belt',
    'price': '$5.00',
    'quantity': '5000',
    'total_purchases': '$25,000.00'
    }]

    transcript = years[0][0]['terms'][0]['classGrades']

    print (transcript)

    document_3 = MailMerge(template_1)
    document_3.merge(**student)
    document_3.merge_rows('classData.subject_name', transcript)
    document_3.write(studentObject['first_name']+' transcript.docx')

def outputDoc(data):

    document = Document('dwight-transcript.docx')

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )



    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')